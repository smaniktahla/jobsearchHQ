"""
Utilities for reading and generating .docx files.
- Extract text from uploaded resume .docx files
- Generate formatted cover letter .docx files matching resume style
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def extract_text_from_docx(file_path: str | Path) -> str:
    """Extract plain text from a .docx file for use as scoring context."""
    doc = Document(str(file_path))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


def generate_cover_letter_docx(
    output_path: str | Path,
    body_text: str,
    candidate_name: str = "Salil Maniktahla",
    candidate_contact: str = "Springfield, VA · 571-215-8218 · salil.maniktahla@gmail.com",
    company: str = "",
    job_title: str = "",
    variant_label: str = "",
):
    """Generate a professional cover letter .docx matching the resume formatting."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Header - Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(candidate_name.upper())
    name_run.bold = True
    name_run.font.size = Pt(15)
    name_run.font.name = "Calibri"
    name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    name_run.font.letter_spacing = Pt(2)

    # Header - Contact
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run(candidate_contact)
    contact_run.font.size = Pt(8.5)
    contact_run.font.name = "Calibri"
    contact_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Divider line
    divider = doc.add_paragraph()
    divider.paragraph_format.space_before = Pt(6)
    divider.paragraph_format.space_after = Pt(12)
    border = divider.paragraph_format
    # Use a bottom border on the paragraph
    from docx.oxml.ns import qn
    pPr = divider._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): "2E5090",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Date
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_before = Pt(6)
    date_para.paragraph_format.space_after = Pt(12)
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(10.5)
    date_run.font.name = "Calibri"

    # Addressee (if company/title known)
    if company or job_title:
        addr_text = ""
        if company:
            addr_text += f"RE: {job_title}" if job_title else f"RE: Position at {company}"
            if job_title and company:
                addr_text = f"RE: {job_title} — {company}"
        addr_para = doc.add_paragraph()
        addr_para.paragraph_format.space_after = Pt(12)
        addr_run = addr_para.add_run(addr_text)
        addr_run.font.size = Pt(10.5)
        addr_run.font.name = "Calibri"
        addr_run.bold = True
        addr_run.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)

    # Body paragraphs
    paragraphs = body_text.strip().split("\n\n")
    for i, para_text in enumerate(paragraphs):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = Pt(15)
        run = p.add_run(para_text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Closing
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(12)
    run = closing.add_run("Best regards,")
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"

    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(4)
    run = sig.add_run(candidate_name)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"

    # Variant watermark (small, bottom)
    if variant_label:
        vp = doc.add_paragraph()
        vp.paragraph_format.space_before = Pt(24)
        vr = vp.add_run(f"[{variant_label}]")
        vr.font.size = Pt(7)
        vr.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        vr.font.name = "Calibri"

    doc.save(str(output_path))
    return str(output_path)
