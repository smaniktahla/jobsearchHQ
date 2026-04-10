"""
Generate .docx resumes matching the Calibri blue-accent template.
Parses structured resume text and produces a properly formatted Word document.
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DATA_DIR = Path("/app/data")
GENERATED_DIR = DATA_DIR / "generated"

# Colors matching the template
DARK = RGBColor(0x1A, 0x1A, 0x1A)
MEDIUM = RGBColor(0x44, 0x44, 0x44)
ACCENT = RGBColor(0x2E, 0x50, 0x90)


def ensure_dirs():
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_border(cell, **kwargs):
    """Set cell border. Unused here but available for table formatting."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        from docx.oxml import OxmlElement
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)


def add_section_heading(doc, text):
    """Add a blue section heading with bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = ACCENT

    # Bottom border via XML
    pPr = p._p.get_or_add_pPr()
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "2E5090")
    pBdr.append(bottom)
    pPr.append(pBdr)

    return p


def add_body_text(doc, text, bold=False, italic=False, color=None, size=9.5):
    """Add a paragraph of body text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1.5)
    p.paragraph_format.space_after = Pt(1.5)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color or DARK
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, size=9.5):
    """Add a bulleted paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0.5)
    p.paragraph_format.space_after = Pt(0.5)
    # Clear default run and add formatted one
    p.clear()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    return p


def add_job_title(doc, text):
    """Add a bold job title."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    return p


def add_job_meta(doc, org, dates=""):
    """Add org name (italic) and right-aligned dates."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0.5)
    p.paragraph_format.space_after = Pt(2)

    run_org = p.add_run(org)
    run_org.italic = True
    run_org.font.name = "Calibri"
    run_org.font.size = Pt(9.5)
    run_org.font.color.rgb = MEDIUM

    if dates:
        # Add tab stop for right-aligned dates
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), "10440")
        tabs.append(tab)
        pPr.append(tabs)

        run_tab = p.add_run("\t")
        run_tab.font.name = "Calibri"
        run_dates = p.add_run(dates)
        run_dates.font.name = "Calibri"
        run_dates.font.size = Pt(9.5)
        run_dates.font.color.rgb = MEDIUM

    return p


def generate_resume_docx(tailored_text: str, job_id: str, company: str = "", title: str = "") -> str:
    """
    Parse tailored resume text and generate a formatted .docx file.
    Returns the path to the generated file.
    """
    ensure_dirs()

    doc = Document()

    # Page setup - US Letter with narrow margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.43)
    section.bottom_margin = Inches(0.43)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(9.5)
    font.color.rgb = DARK

    # Parse the text and build the document
    lines = tailored_text.strip().split("\n")
    i = 0
    in_header = True

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Detect section headers (all caps lines, or lines ending with colon that look like headers)
        is_section = (
            line.isupper() and len(line) > 3 and not line.startswith("•") and not line.startswith("-")
        ) or (
            line.upper() == line and "&" in line and len(line) > 5
        )

        # Also detect headers like "SUMMARY", "CORE CAPABILITIES", etc.
        known_sections = [
            "SUMMARY", "CORE CAPABILITIES", "PROFESSIONAL EXPERIENCE",
            "EARLIER EXPERIENCE", "TECHNICAL SKILLS", "CERTIFICATIONS",
            "EDUCATION", "CERTIFICATIONS & EDUCATION", "SELECTED INNOVATION",
            "SELECTED SYSTEMS & INNOVATION", "PROJECT HIGHLIGHTS", "KEY SKILLS",
        ]
        if line.upper().rstrip(":") in known_sections:
            is_section = True

        # Header block (name + contact)
        if in_header and i < 3:
            if i == 0:
                # Name
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(line.upper())
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(15)
                run.font.color.rgb = DARK
                run.font.cs_bold = True
                # Letter spacing
                rPr = run._r.get_or_add_rPr()
                from docx.oxml import OxmlElement
                spacing = OxmlElement("w:spacing")
                spacing.set(qn("w:val"), "50")
                rPr.append(spacing)
            elif "·" in line or "@" in line or "linkedin" in line.lower():
                # Contact line
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(1.5)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(line)
                run.font.name = "Calibri"
                run.font.size = Pt(8.5)
                run.font.color.rgb = MEDIUM
                in_header = False
            else:
                in_header = False
                continue
            i += 1
            continue

        in_header = False

        # Section heading
        if is_section:
            add_section_heading(doc, line.rstrip(":"))
            i += 1
            continue

        # Bullet point
        if line.startswith("•") or line.startswith("-") or line.startswith("*"):
            text = line.lstrip("•-* ").strip()
            add_bullet(doc, text)
            i += 1
            continue

        # Sub-headers (bold lines like "BI Senior Manager (Aug 2012 – Jul 2016)")
        if re.match(r"^(BI |Senior |Operations |IT |Co-Founder|Product )", line) and "(" in line:
            add_body_text(doc, line, bold=True, size=9.5)
            i += 1
            continue

        # Job title detection (bold, followed by org on next line)
        next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
        looks_like_title = (
            not line.startswith("•") and
            len(line) < 80 and
            (next_line and ("–" in next_line or "—" in next_line or "·" in next_line or re.match(r"^[A-Z]", next_line)))
        )

        # Check if this is a job entry (title + org/dates pattern)
        date_pattern = r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Present|\d{4})"
        if re.search(date_pattern, next_line) and looks_like_title:
            add_job_title(doc, line)
            # Parse next line for org and dates
            i += 1
            meta_line = lines[i].strip() if i < len(lines) else ""
            # Try to split on tab or multiple spaces for right-aligned dates
            if "\t" in meta_line:
                parts = meta_line.split("\t")
                add_job_meta(doc, parts[0].strip(), parts[-1].strip() if len(parts) > 1 else "")
            elif re.search(date_pattern, meta_line):
                # Try to find date range at end
                match = re.search(r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Present)[\w\s,–\-]*\d{4}.*$)", meta_line)
                if match:
                    org = meta_line[:match.start()].rstrip(" ·–—,")
                    dates = match.group(1).strip()
                    add_job_meta(doc, org, dates)
                else:
                    add_job_meta(doc, meta_line)
            else:
                add_job_meta(doc, meta_line)
            i += 1
            continue

        # Skill lines (Category: items)
        if ":" in line and len(line.split(":")[0]) < 25:
            parts = line.split(":", 1)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run_label = p.add_run(parts[0] + ": ")
            run_label.bold = True
            run_label.font.name = "Calibri"
            run_label.font.size = Pt(9.5)
            run_label.font.color.rgb = DARK
            run_val = p.add_run(parts[1].strip())
            run_val.font.name = "Calibri"
            run_val.font.size = Pt(9.5)
            run_val.font.color.rgb = DARK
            i += 1
            continue

        # Default: body text
        add_body_text(doc, line)
        i += 1

    # Save
    safe_company = re.sub(r"[^\w\s-]", "", company or "job")[:30].strip()
    filename = f"Resume_Tailored_{safe_company}_{job_id}.docx"
    filepath = GENERATED_DIR / filename
    doc.save(str(filepath))

    return str(filepath)


def generate_cover_letter_docx(
    content: str,
    variant: str,
    job_id: str,
    company: str = "",
    title: str = "",
) -> str:
    """
    Generate a professional cover letter .docx.
    Clean business letter format matching the resume's Calibri style.
    """
    ensure_dirs()

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK

    # Header — name and contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Salil Maniktahla")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Springfield, VA  \u00B7  571-215-8218  \u00B7  salil.maniktahla@gmail.com")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MEDIUM

    # Thin rule
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "2E5090")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Date
    from datetime import datetime
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = MEDIUM

    # Body paragraphs
    paragraphs = content.strip().split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        # Handle single newlines within a paragraph as continuous text
        para_text = para_text.replace("\n", " ")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = Pt(15)
        run = p.add_run(para_text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = DARK

    # Save
    safe_company = re.sub(r"[^\w\s-]", "", company or "job")[:30].strip()
    safe_variant = re.sub(r"[^\w]", "", variant)[:15]
    filename = f"CoverLetter_{safe_variant}_{safe_company}_{job_id}.docx"
    filepath = GENERATED_DIR / filename
    doc.save(str(filepath))

    return str(filepath)
